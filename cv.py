from docx import Document

# Create a new Document
doc = Document()

# Add Name and Contact Info
doc.add_heading("Johnpaul Kibet", level=0)
doc.add_paragraph("Nairobi, Kenya | johnpaulkibbet3@gmail.com | 0712492105")
doc.add_paragraph("LinkedIn | GitHub")

# Add Professional Summary
doc.add_heading("Professional Summary", level=1)
doc.add_paragraph(
    "Motivated and detail-oriented computer science enthusiast with a strong foundation in programming, "
    "data structures, and web technologies. Proficient in C, C++, Python, SQL, HTML, and CSS. Passionate "
    "about solving real-world problems and learning new technologies. Seeking opportunities to contribute "
    "to impactful projects while expanding technical and professional skills."
)

# Add Education
doc.add_heading("Education", level=1)
doc.add_paragraph("BSc in Software Engineering, Multimedia University of Kenya")
doc.add_paragraph("August 2023 – Present")

# Add Skills
doc.add_heading("Skills", level=1)
doc.add_paragraph("Technical: C, C++, Python, HTML, CSS, SQL, Git, Data Structures, Algorithms, OOP")
doc.add_paragraph("Soft Skills: Problem-Solving, Communication, Teamwork, Adaptability, Attention to Detail")

# Add Projects
doc.add_heading("Projects", level=1)
doc.add_paragraph("Mbogi Coop", style='List Bullet')
doc.add_paragraph(
    "A web app that helps youth save collectively and support each other financially.\n"
    "Tech Stack: HTML, CSS, JavaScript, PHP\n"
    "Role: Lead Idea Pitcher & Contributor", style='Normal'
)

# Add Experience
doc.add_heading("Experience", level=1)
doc.add_paragraph("Open to internship opportunities to apply and expand my skills.")

# Add Certifications
doc.add_heading("Certifications", level=1)
doc.add_paragraph("Introduction to Cyber Security | Cisco | February 2025")

# Save the document
doc.save("Johnpaul_Kibet_Resume.docx")